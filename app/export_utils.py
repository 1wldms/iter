import os
from io import BytesIO
from datetime import datetime

from reportlab.lib.pagesizes import A4
from reportlab.lib.units import mm
from reportlab.lib import colors
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak

from docx import Document
from docx.shared import Pt

# ── 한글 폰트 등록 (PDF) ──────────────────────────────────
# app/fonts/ 폴더에 NotoSansKR-Regular.ttf, NotoSansKR-Bold.ttf 를 넣어주세요.
# 다운로드: https://fonts.google.com/noto/specimen/Noto+Sans+KR
FONT_PATH = os.path.join(os.path.dirname(__file__), "fonts", "static", "NotoSansKR-Regular.ttf")
FONT_BOLD_PATH = os.path.join(os.path.dirname(__file__), "fonts", "static", "NotoSansKR-Bold.ttf")
FONT_NAME = "NotoSansKR"
FONT_NAME_BOLD = "NotoSansKR-Bold"

_fonts_registered = False


def _register_fonts():
    global _fonts_registered
    if _fonts_registered:
        return
    if not os.path.exists(FONT_PATH):
        raise FileNotFoundError(
            f"한글 폰트 파일이 없어요: {FONT_PATH}\n"
            "https://fonts.google.com/noto/specimen/Noto+Sans+KR 에서 다운로드해서 app/fonts/ 에 넣어주세요."
        )
    pdfmetrics.registerFont(TTFont(FONT_NAME, FONT_PATH))
    if os.path.exists(FONT_BOLD_PATH):
        pdfmetrics.registerFont(TTFont(FONT_NAME_BOLD, FONT_BOLD_PATH))
    else:
        pdfmetrics.registerFont(TTFont(FONT_NAME_BOLD, FONT_PATH))
    _fonts_registered = True


FIELD_LABELS = {
    "role": "역할",
    "background": "배경",
    "action": "행동",
    "result": "결과",
    "learned": "배운 점",
    "reflection": "느낀 점",
    "memo": "메모",
}

FIELD_ORDER = ["role", "background", "action", "result", "learned", "reflection", "memo"]


def _format_date_range(exp):
    start = exp.get("start_date")
    end = exp.get("end_date")
    if not start and not end:
        return ""
    start_s = start[:7] if start else "?"
    end_s = end[:7] if end else "현재"
    return f"{start_s} ~ {end_s}"


# ── PDF 생성 ──────────────────────────────────────────────
def _get_pdf_styles():
    _register_fonts()
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "TitleKR", parent=styles["Title"], fontName=FONT_NAME_BOLD, fontSize=18, spaceAfter=4,
    )
    meta_style = ParagraphStyle(
        "MetaKR", parent=styles["Normal"], fontName=FONT_NAME, fontSize=9,
        textColor=colors.grey, spaceAfter=12,
    )
    label_style = ParagraphStyle(
        "LabelKR", parent=styles["Heading3"], fontName=FONT_NAME_BOLD, fontSize=11,
        textColor=colors.HexColor("#4B4033"), spaceBefore=10, spaceAfter=4,
    )
    body_style = ParagraphStyle(
        "BodyKR", parent=styles["Normal"], fontName=FONT_NAME, fontSize=10.5,
        leading=16, spaceAfter=4,
    )
    keyword_style = ParagraphStyle(
        "KeywordKR", parent=styles["Normal"], fontName=FONT_NAME, fontSize=9,
        textColor=colors.HexColor("#8A7A5C"),
    )
    return title_style, meta_style, label_style, body_style, keyword_style


def _escape(text: str) -> str:
    return (
        text.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
        .replace("\n", "<br/>")
    )


def _experience_flowables(exp, styles, add_pagebreak_after=False):
    title_style, meta_style, label_style, body_style, keyword_style = styles
    flow = []

    title = exp.get("title") or "(제목 없음)"
    flow.append(Paragraph(_escape(title), title_style))

    date_range = _format_date_range(exp)
    if date_range:
        flow.append(Paragraph(date_range, meta_style))
    else:
        flow.append(Spacer(1, 8))

    keywords = exp.get("keywords") or []
    if keywords:
        flow.append(Paragraph(" · ".join(f"#{k}" for k in keywords), keyword_style))
        flow.append(Spacer(1, 6))

    for field in FIELD_ORDER:
        value = (exp.get(field) or "").strip()
        if not value:
            continue
        flow.append(Paragraph(FIELD_LABELS[field], label_style))
        flow.append(Paragraph(_escape(value), body_style))

    if add_pagebreak_after:
        flow.append(PageBreak())

    return flow


def generate_experience_pdf(experience: dict) -> BytesIO:
    buf = BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=A4,
        topMargin=20 * mm, bottomMargin=20 * mm, leftMargin=20 * mm, rightMargin=20 * mm,
    )
    styles = _get_pdf_styles()
    doc.build(_experience_flowables(experience, styles))
    buf.seek(0)
    return buf


def generate_experiences_list_pdf(experiences: list, owner_name: str = "") -> BytesIO:
    buf = BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=A4,
        topMargin=20 * mm, bottomMargin=20 * mm, leftMargin=20 * mm, rightMargin=20 * mm,
    )
    styles = _get_pdf_styles()
    title_style, meta_style, label_style, body_style, keyword_style = styles

    flow = []
    header = f"{owner_name}의 경험 아카이브" if owner_name else "경험 아카이브"
    flow.append(Paragraph(header, title_style))
    flow.append(Paragraph(f"생성일: {datetime.now().strftime('%Y-%m-%d')} · 총 {len(experiences)}개", meta_style))
    flow.append(Spacer(1, 10))

    for i, exp in enumerate(experiences):
        is_last = i == len(experiences) - 1
        flow.extend(_experience_flowables(exp, styles, add_pagebreak_after=not is_last))

    doc.build(flow)
    buf.seek(0)
    return buf


# ── Word(.docx) 생성 ──────────────────────────────────────
def _add_experience_to_doc(doc: Document, exp: dict):
    title = exp.get("title") or "(제목 없음)"
    doc.add_heading(title, level=1)

    date_range = _format_date_range(exp)
    if date_range:
        p = doc.add_paragraph(date_range)
        p.runs[0].font.size = Pt(9)

    keywords = exp.get("keywords") or []
    if keywords:
        doc.add_paragraph(" · ".join(f"#{k}" for k in keywords))

    for field in FIELD_ORDER:
        value = (exp.get(field) or "").strip()
        if not value:
            continue
        doc.add_heading(FIELD_LABELS[field], level=3)
        for line in value.split("\n"):
            doc.add_paragraph(line if line.strip() else "")


def generate_experience_docx(experience: dict) -> BytesIO:
    doc = Document()
    _add_experience_to_doc(doc, experience)
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def generate_experiences_list_docx(experiences: list, owner_name: str = "") -> BytesIO:
    doc = Document()
    header = f"{owner_name}의 경험 아카이브" if owner_name else "경험 아카이브"
    doc.add_heading(header, level=0)
    doc.add_paragraph(f"생성일: {datetime.now().strftime('%Y-%m-%d')} · 총 {len(experiences)}개")

    for i, exp in enumerate(experiences):
        if i > 0:
            doc.add_page_break()
        _add_experience_to_doc(doc, exp)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf